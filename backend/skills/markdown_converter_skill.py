"""
Pandoc + Mermaid + CSV 完整工作流
将 Markdown 转换为 PDF/Word/HTML/Excel，支持 Mermaid 图表和 CSV 表格

依赖安装:
    pip install pypandoc python-docx openpyxl pdfplumber markdown mermaiddraft pyyaml

Mermaid 渲染选项 (需要安装):
    - @mermaid-js/mermaid-cli (npm install -g @mermaid-js/mermaid-cli)
    - 或使用 mermaid.ink 在线服务
"""

import base64
import io
import json
import logging
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from enum import Enum

import httpx

logger = logging.getLogger(__name__)


class OutputFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    EXCEL = "xlsx"
    PPTX = "pptx"
    MARKDOWN = "markdown"


@dataclass
class ConversionResult:
    success: bool
    file_path: Optional[str] = None
    content: Optional[bytes] = None
    error: Optional[str] = None
    warnings: List[str] = None
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


class MermaidRenderer:
    """Mermaid 图表渲染器"""
    
    def __init__(self, use_online: bool = True):
        """
        初始化 Mermaid 渲染器
        
        Args:
            use_online: 是否使用 mermaid.ink 在线服务 (否则需要本地 mmdc)
        """
        self.use_online = use_online
        self.mmdc_path = self._find_mmdc()
    
    def _find_mmdc(self) -> Optional[str]:
        """查找本地 mermaid-cli"""
        import shutil
        path = shutil.which('mmdc')
        if path:
            return path
        
        # 常见路径
        common_paths = [
            r"C:\Program Files\nodejs\mmdc.cmd",
            r"C:\Users\AppData\Roaming\npm\mmdc.cmd",
        ]
        for p in common_paths:
            if os.path.exists(p):
                return p
        return None
    
    async def render_to_svg(self, mermaid_code: str) -> Optional[str]:
        """将 Mermaid 代码渲染为 SVG"""
        if self.use_online:
            return await self._render_online(mermaid_code, 'svg')
        else:
            return await self._render_local(mermaid_code, 'svg')
    
    async def render_to_png(self, mermaid_code: str) -> Optional[str]:
        """将 Mermaid 代码渲染为 PNG"""
        if self.use_online:
            return await self._render_online(mermaid_code, 'png')
        else:
            return await self._render_local(mermaid_code, 'png')
    
    async def _render_online(self, mermaid_code: str, format: str) -> Optional[str]:
        """使用 mermaid.ink 在线服务渲染"""
        try:
            # 清理代码
            code = mermaid_code.strip()
            encoded = base64.urlsafe_b64encode(code.encode()).decode()
            
            url = f"https://mermaid.ink/{format}/{encoded}"
            
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url)
                
                if response.status_code == 200:
                    if format == 'svg':
                        return response.text
                    else:
                        return base64.b64encode(response.content).decode()
            return None
        except Exception as e:
            logger.warning(f"Mermaid online render failed: {e}")
            return None
    
    async def _render_local(self, mermaid_code: str, format: str) -> Optional[str]:
        """使用本地 mmdc 渲染"""
        if not self.mmdc_path:
            logger.warning("mmdc not found, falling back to online")
            return await self._render_online(mermaid_code, format)
        
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                input_file = Path(tmpdir) / "input.mmd"
                output_file = Path(tmpdir) / f"output.{format}"
                
                input_file.write_text(mermaid_code, encoding='utf-8')
                
                cmd = [
                    self.mmdc_path,
                    '-i', str(input_file),
                    '-o', str(output_file),
                    '-b', 'transparent',
                    '-w', '1200',
                    '-H', '800'
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0 and output_file.exists():
                    if format == 'svg':
                        return output_file.read_text(encoding='utf-8')
                    else:
                        return base64.b64encode(output_file.read_bytes()).decode()
            return None
        except Exception as e:
            logger.warning(f"Mermaid local render failed: {e}")
            return None


class CSVTableExtractor:
    """CSV 表格智能提取器"""
    
    def __init__(self):
        self.pdfplumber_available = self._check_pdfplumber()
    
    def _check_pdfplumber(self) -> bool:
        try:
            import pdfplumber
            return True
        except ImportError:
            return False
    
    async def extract_tables_from_pdf(self, pdf_path: str) -> List[List[List[str]]]:
        """从 PDF 提取表格数据"""
        if not self.pdfplumber_available:
            return self._fallback_extract(pdf_path)
        
        try:
            import pdfplumber
            
            tables = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
            
            return tables
        except Exception as e:
            logger.warning(f"PDF table extraction failed: {e}")
            return self._fallback_extract(pdf_path)
    
    def _fallback_extract(self, pdf_path: str) -> List[List[List[str]]]:
        """回退方案：简单文本提取"""
        try:
            with open(pdf_path, 'rb') as f:
                content = f.read()
                text = content.decode('utf-8', errors='ignore')
            
            # 简单按行分割
            lines = text.split('\n')[:100]
            return [[line.strip().split('\t')] for line in lines if line.strip()]
        except Exception as e:
            logger.error(f"Fallback extraction failed: {e}")
            return []
    
    def tables_to_csv(self, tables: List[List[List[str]]]) -> str:
        """将表格转换为 CSV 格式"""
        import csv
        from io import StringIO
        
        output = StringIO()
        writer = csv.writer(output)
        
        for i, table in enumerate(tables):
            if i > 0:
                writer.writerow([])  # 空行分隔
            for row in table:
                writer.writerow(row)
        
        return output.getvalue()


class MarkdownPreprocessor:
    """Markdown 预处理器 - 处理 Mermaid 图表和 CSV 表格"""
    
    def __init__(self):
        self.mermaid_renderer = MermaidRenderer()
    
    async def preprocess(self, markdown_content: str, render_mermaid: bool = True) -> str:
        """
        预处理 Markdown 内容
        
        Args:
            markdown_content: 原始 Markdown 内容
            render_mermaid: 是否渲染 Mermaid 图表
            
        Returns:
            处理后的 Markdown (Mermaid 块被替换为图片)
        """
        if not render_mermaid:
            return markdown_content
        
        # 匹配 Mermaid 代码块
        pattern = r'```mermaid\s*\n(.*?)```'
        
        async def replace_mermaid(match):
            mermaid_code = match.group(1)
            svg = await self.mermaid_renderer.render_to_svg(mermaid_code)
            
            if svg:
                # 返回 HTML img 标签
                return f'\n![Mermaid Diagram](data:image/svg+xml;base64,{base64.b64encode(svg.encode()).decode()})\n'
            else:
                # 渲染失败，保留原始代码
                return match.group(0)
        
        # 使用同步方式处理（实际应该用 asyncio 但 re.sub 不支持 async）
        result = re.sub(pattern, replace_mermaid, markdown_content, flags=re.DOTALL)
        
        return result


class MarkdownConverter:
    """Markdown 文档转换器 - 直接使用 LibreOffice，不依赖 pandoc"""

    # PDF 导出的 10 种主题色
    THEMES = {
        'warm-academic': {'primary':'#C17B4B', 'secondary':'#8B5E3C', 'accent':'#E8D5C4'},
        'classic-thesis': {'primary':'#8B4513', 'secondary':'#6B3410', 'accent':'#F5E6D3'},
        'tufte': {'primary':'#8B0000', 'secondary':'#4A4A4A', 'accent':'#F0F0F0'},
        'ieee-journal': {'primary':'#1B3A5C', 'secondary':'#2F5496', 'accent':'#E8EEF4'},
        'elegant-book': {'primary':'#6B4226', 'secondary':'#8B6914', 'accent':'#FAF0E6'},
        'chinese-red': {'primary':'#CC2936', 'secondary':'#8B1A1A', 'accent':'#FFF8F0'},
        'ink-wash': {'primary':'#2D2D2D', 'secondary':'#595959', 'accent':'#F8F8F8'},
        'github-light': {'primary':'#0366D6', 'secondary':'#586069', 'accent':'#F6F8FA'},
        'nord-frost': {'primary':'#5E81AC', 'secondary':'#81A1C1', 'accent':'#ECEFF4'},
        'ocean-breeze': {'primary':'#00897B', 'secondary':'#00695C', 'accent':'#E0F2F1'},
    }

    def __init__(self):
        self.preprocessor = MarkdownPreprocessor()
        self.lo_path = self._find_libreoffice_path()
        logger.info(f"[MarkdownConverter] LibreOffice path: {self.lo_path}")

    def _find_libreoffice_path(self) -> Optional[str]:
        """查找 LibreOffice 可执行文件"""
        import shutil
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for p in paths:
            if Path(p).exists():
                return p
        soffice = shutil.which("soffice")
        if soffice:
            return soffice
        return None

    async def convert(
        self,
        input_content: str,
        input_format: str,
        output_format: OutputFormat,
        render_mermaid: bool = True,
        extract_csv: bool = False,
        theme: str = 'chinese-red'
    ) -> ConversionResult:
        """
        转换文档: Markdown → PDF/DOCX/HTML/Excel

        链路: LibreOffice → (Python 原生兜底)
        """
        try:
            # 1. 预处理 Markdown (Mermaid 图表)
            if input_format == 'markdown':
                processed_content = await self.preprocessor.preprocess(
                    input_content,
                    render_mermaid=render_mermaid
                )
            else:
                processed_content = input_content

            # 2. Excel/CSV 特殊处理
            if output_format == OutputFormat.EXCEL and extract_csv:
                return await self._convert_to_excel_with_csv(processed_content)

            # 3. LibreOffice 转换
            return await self._convert_with_libreoffice(processed_content, output_format, theme)

        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            return ConversionResult(success=False, error=str(e))

    async def _convert_with_libreoffice(
        self,
        content: str,
        output_format: OutputFormat,
        theme: str = 'chinese-red'
    ) -> ConversionResult:
        """使用 LibreOffice 转换"""
        if not self.lo_path:
            # 直接用 Python 原生 fallback
            return await self._fallback_reportlab(content, output_format, theme)

        with tempfile.TemporaryDirectory() as tmpdir:
            tmppath = Path(tmpdir)

            # 写入临时文件
            input_file = tmppath / "input.md"
            input_file.write_text(content, encoding='utf-8')

            output_ext = output_format.value
            output_file = tmppath / f"output.{output_ext}"

            # Markdown → HTML (简单转换，再由 LibreOffice 转目标格式)
            html_content = self._markdown_to_html(content)
            html_file = tmppath / "input.html"
            html_file.write_text(html_content, encoding='utf-8')

            # LibreOffice 能直接读 HTML
            import uuid
            user_install = tmppath / f".config_{uuid.uuid4().hex[:8]}"
            os.environ['HOME'] = str(tmppath)
            cmd = [
                self.lo_path,
                "--headless", "--norestore", "--nofirststartwizard",
                "-env:UserInstallation=file:///" + str(user_install).replace('\\', '/'),
                "--convert-to", output_ext,
                "--outdir", str(tmppath),
                str(html_file)
            ]

            logger.info(f"[LibreOffice] Executing: {' '.join(cmd)}")

            try:
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=120
                )

                if result.returncode == 0 and output_file.exists():
                    return ConversionResult(
                        success=True,
                        file_path=str(output_file),
                        content=output_file.read_bytes()
                    )
                else:
                    logger.warning(f"[LibreOffice] failed: {result.stderr}, trying reportlab")
                    return await self._fallback_reportlab(content, output_format, theme)

            except subprocess.TimeoutExpired:
                return ConversionResult(success=False, error="LibreOffice conversion timeout")

    def _markdown_to_html(self, md_content: str) -> str:
        """简单的 Markdown 转 HTML（不依赖外部库）"""
        import html as html_lib
        lines_text = md_content.split('\n')
        html_parts = []
        in_code = False
        in_list = False

        for line_text in lines_text:
            if line_text.strip().startswith('```'):
                if not in_code:
                    html_parts.append('<pre><code>')
                    in_code = True
                else:
                    html_parts.append('</code></pre>')
                    in_code = False
                continue

            if in_code:
                html_parts.append(html_lib.escape(line_text))
                continue

            if line_text.startswith('# '):
                html_parts.append(f'<h1>{html_lib.escape(line_text[2:])}</h1>')
            elif line_text.startswith('## '):
                html_parts.append(f'<h2>{html_lib.escape(line_text[3:])}</h2>')
            elif line_text.startswith('### '):
                html_parts.append(f'<h3>{html_lib.escape(line_text[4:])}</h3>')
            elif line_text.strip().startswith('- ') or line_text.strip().startswith('* '):
                if not in_list:
                    html_parts.append('<ul>')
                    in_list = True
                html_parts.append(f'<li>{html_lib.escape(line_text.strip()[2:])}</li>')
            else:
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                if line_text.strip():
                    html_parts.append(f'<p>{html_lib.escape(line_text)}</p>')

        if in_list:
            html_parts.append('</ul>')

        return '<!DOCTYPE html>\n<html>\n<head>\n<meta charset="utf-8">\n'                '<style>body{font-family:"Noto Sans SC",sans-serif;padding:20px;} '                'h1,h2,h3{color:#333;}code{background:#f5f5f5;padding:2px 6px;} '                'pre{background:#f5f5f5;padding:15px;overflow-x:auto;}</style>\n'                '</head>\n<body>\n' + '\n'.join(html_parts) + '\n</body>\n</html>'

    async def _fallback_reportlab(
            self,
            content: str,
            output_format: OutputFormat,
            theme: str = 'chinese-red'
    ) -> ConversionResult:
        """最后 fallback：使用 Python 直接生成目标格式"""
        try:
            tc = self.THEMES.get(theme, self.THEMES['chinese-red'])
            primary_color = tc['primary']
            accent_color = tc['accent']

            # HTML: 用 markdown 库直接转，带主题色
            if output_format == OutputFormat.HTML:
                import markdown as md_lib
                html_body = md_lib.markdown(content, extensions=['fenced_code', 'tables'])
                styled_html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; color: #333; background: {accent_color}; }}
a {{ color: {primary_color}; }}
h1, h2, h3 {{ color: {primary_color}; }}
pre {{ background: #f5f5f5; padding: 12px; border-radius: 6px; overflow-x: auto; border-left: 4px solid {primary_color}; }}
code {{ background: #f0f0f0; padding: 2px 5px; border-radius: 3px; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background: {primary_color}; color: white; }}
</style></head><body>{html_body}</body></html>"""
                return ConversionResult(success=True, content=styled_html.encode('utf-8'))

            # DOCX: 用 python-docx 创建 Word 文档
            if output_format == OutputFormat.DOCX:
                from docx import Document
                from docx.shared import Pt, Inches
                doc = Document()
                lines = content.split('\n')
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith('# ') or stripped.startswith('## '):
                        level = 1 if stripped.startswith('# ') else 2
                        text = stripped.lstrip('#').strip()
                        h = doc.add_heading(text, level=level)
                    elif stripped.startswith('- ') or stripped.startswith('* '):
                        doc.add_paragraph(stripped[2:], style='List Bullet')
                    elif stripped:
                        p = doc.add_paragraph(stripped)
                    else:
                        doc.add_paragraph()
                buffer = io.BytesIO()
                doc.save(buffer)
                return ConversionResult(success=True, content=buffer.getvalue())

            # PDF（原有 reportlab 逻辑）
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from io import BytesIO

            # 查找中文字体
            font_paths = [
                Path(os.environ.get('WINDIR', 'C:\\Windows')) / 'Fonts' / 'msyh.ttc',
                Path.home() / 'AppData/Local/Microsoft/Windows/Fonts' / 'NotoSansSC-Regular.ttf',
                Path(__file__).parent.parent.parent / 'public' / 'fonts' / 'NotoSansSC-Regular.ttf',
            ]
            font_registered = False
            for fp in font_paths:
                if fp.exists():
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', str(fp), 'Identity-H'))
                        font_registered = True
                        break
                    except Exception:
                        continue

            if not font_registered:
                return ConversionResult(success=False, error="No Chinese font available")

            paragraphs = self._extract_text_paragraphs(content)

            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                leftMargin=20*mm, rightMargin=20*mm,
                topMargin=20*mm, bottomMargin=20*mm
            )

            style = ParagraphStyle('Chinese', fontName='ChineseFont', fontSize=11, leading=16)
            title_style = ParagraphStyle('Title', fontName='ChineseFont', fontSize=18, leading=24, spaceAfter=12)

            story = []
            for para in paragraphs:
                if para.get('type') == 'h1':
                    story.append(Paragraph(para['text'], title_style))
                    story.append(Spacer(1, 6*mm))
                elif para.get('type') == 'h2':
                    story.append(Paragraph(para['text'], ParagraphStyle('H2', fontName='ChineseFont', fontSize=14, leading=20, spaceAfter=6)))
                    story.append(Spacer(1, 3*mm))
                elif para.get('type') == 'text':
                    story.append(Paragraph(para['text'], style))
                    story.append(Spacer(1, 3*mm))

            doc.build(story)
            return ConversionResult(success=True, content=buffer.getvalue())

        except ImportError as e:
            return ConversionResult(success=False, error=f"required library not available: {e}")
        except Exception as e:
            logger.error(f"Fallback conversion failed: {e}")
            return ConversionResult(success=False, error=str(e))

    def _extract_text_paragraphs(self, content: str) -> list:
        """从 markdown 提取纯文本段落"""
        import html as html_lib
        lines_text = content.split('\n')
        paragraphs = []
        in_code = False

        for line_text in lines_text:
            if line_text.strip().startswith('```'):
                in_code = not in_code
                continue
            if in_code:
                continue

            if line_text.startswith('# '):
                paragraphs.append({'type': 'h1', 'text': html_lib.escape(line_text[2:])})
            elif line_text.startswith('## '):
                paragraphs.append({'type': 'h2', 'text': html_lib.escape(line_text[3:])})
            elif line_text.startswith('### '):
                paragraphs.append({'type': 'h3', 'text': html_lib.escape(line_text[4:])})
            elif line_text.strip().startswith('- ') or line_text.strip().startswith('* '):
                paragraphs.append({'type': 'text', 'text': f"• {html_lib.escape(line_text.strip()[2:])}"})
            elif line_text.strip():
                paragraphs.append({'type': 'text', 'text': html_lib.escape(line_text)})

        return paragraphs

    async def _convert_to_excel_with_csv(self, content: str) -> ConversionResult:
        """将 CSV 数据转换为 Excel"""
        try:
            from openpyxl import Workbook
            csv_pattern = r'```csv\s*\n(.*?)```'
            matches = re.findall(csv_pattern, content, re.DOTALL)

            if not matches:
                return ConversionResult(success=False, error="No CSV data found in content")

            wb = Workbook()
            ws = wb.active
            ws.title = "Data"

            import csv
            from io import StringIO
            for i, csv_data in enumerate(matches):
                if i > 0:
                    ws.append([])
                reader = csv.reader(StringIO(csv_data))
                for row in reader:
                    ws.append(row)

            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as f:
                wb.save(f.name)
                return ConversionResult(
                    success=True,
                    file_path=f.name,
                    content=Path(f.name).read_bytes()
                )

        except Exception as e:
            logger.error(f"Excel conversion failed: {e}")
            return ConversionResult(success=False, error=str(e))

# 便捷函数
async def markdown_to_pdf(
    content: str,
    render_mermaid: bool = True
) -> ConversionResult:
    """Markdown 转 PDF"""
    converter = MarkdownConverter()
    return await converter.convert(
        content, 'markdown', OutputFormat.PDF, render_mermaid
    )


async def markdown_to_docx(
    content: str,
    render_mermaid: bool = True
) -> ConversionResult:
    """Markdown 转 Word"""
    converter = MarkdownConverter()
    return await converter.convert(
        content, 'markdown', OutputFormat.DOCX, render_mermaid
    )


async def markdown_to_html(
    content: str,
    render_mermaid: bool = True
) -> ConversionResult:
    """Markdown 转 HTML"""
    converter = MarkdownConverter()
    return await converter.convert(
        content, 'markdown', OutputFormat.HTML, render_mermaid
    )


async def markdown_to_excel(
    content: str,
    extract_csv: bool = True
) -> ConversionResult:
    """Markdown 转 Excel (提取 CSV 表格)"""
    converter = MarkdownConverter()
    return await converter.convert(
        content, 'markdown', OutputFormat.EXCEL, 
        render_mermaid=False, extract_csv=extract_csv
    )


async def pdf_tables_to_csv(pdf_path: str) -> str:
    """从 PDF 提取表格为 CSV"""
    extractor = CSVTableExtractor()
    tables = await extractor.extract_tables_from_pdf(pdf_path)
    return extractor.tables_to_csv(tables)


# 示例用法
async def demo():
    """演示完整工作流"""
    
    # 示例 Markdown 内容
    md_content = """
# 项目报告

## 流程图

```mermaid
graph TD
    A[开始] --> B{判断}
    B -->|是| C[处理]
    B -->|否| D[结束]
    C --> D
```

## 数据表格

```csv
日期,销售额,成本,利润
2024-01,10000,6000,4000
2024-02,12000,7000,5000
2024-03,15000,8000,7000
```

## 结论

报告完成。
"""
    
    # 转换为不同格式
    print("Converting to PDF...")
    result_pdf = await markdown_to_pdf(md_content)
    print(f"PDF: {result_pdf.success}")
    
    print("Converting to DOCX...")
    result_docx = await markdown_to_docx(md_content)
    print(f"DOCX: {result_docx.success}")
    
    print("Converting to HTML...")
    result_html = await markdown_to_html(md_content)
    print(f"HTML: {result_html.success}")
    
    print("Converting to Excel...")
    result_xlsx = await markdown_to_excel(md_content)
    print(f"Excel: {result_xlsx.success}")


if __name__ == "__main__":
    import asyncio
    asyncio.run(demo())