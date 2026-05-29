"""
PDF 四色卡片处理器
用于从 PDF 提取内容并生成四色知识卡片
"""

import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# 如果 pypdf 和 pdfplumber 都不可用，标记为不可用
if not PYPDF_AVAILABLE and not PDFPLUMBER_AVAILABLE:
    PDF_AVAILABLE = False
else:
    PDF_AVAILABLE = True

logger = logging.getLogger(__name__)

REPORTLAB_AVAILABLE = True
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
except ImportError:
    REPORTLAB_AVAILABLE = False


# ============ 中文字体注册（单例）============
_FONT_REGISTERED = False
_CHINESE_FONT_NAME = 'Helvetica'


def _get_chinese_font_path():
    """查找中文字体路径"""
    base_dirs = [
        Path(__file__).parent.parent.parent,  # 项目根目录
    ]
    font_names = ["NotoSansSC-Regular.ttf", "SimHei.ttf"]
    for base in base_dirs:
        for subdir in ["public/fonts", "fonts"]:
            for fname in font_names:
                path = base / subdir / fname
                if path.exists():
                    return str(path)
    return None


def _ensure_chinese_font():
    """确保中文字体已注册到 reportlab"""
    global _FONT_REGISTERED, _CHINESE_FONT_NAME
    if _FONT_REGISTERED:
        return _CHINESE_FONT_NAME
    _FONT_REGISTERED = True

    font_path = _get_chinese_font_path()
    if font_path:
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            # 注册常规体
            pdfmetrics.registerFont(TTFont('NotoSansSC', font_path, 'Identity-H'))
            # 注册粗体（如果存在单独文件，否则复用同一文件，reportlab 会做向量粗化）
            bold_path = font_path.replace('Regular', 'Bold')
            if not Path(bold_path).exists():
                bold_path = font_path  # 用同一文件，reportlab 自动加粗
            pdfmetrics.registerFont(TTFont('NotoSansSC-Bold', bold_path, 'Identity-H'))
            _CHINESE_FONT_NAME = 'NotoSansSC'
            logger.info(f"[FourColorPDF] 中文字体注册成功: {font_path}")
            return 'NotoSansSC'
        except Exception as e:
            logger.warning(f"[FourColorPDF] 中文字体注册失败: {e}")

    logger.warning("[FourColorPDF] 未找到中文字体，使用 Helvetica（可能显示乱码）")
    return 'Helvetica'


class PDFourColorProcessor:
    """PDF 四色卡片处理器"""
    
    # 四色卡片类型定义
    CARD_TYPES = {
        "fact": {
            "name": "事实",
            "color": "blue",
            "keywords": ["数据", "统计", "数量", "金额", "比例", "增长", "下降", "达到", "超过", "约", "共"],
            "description": "客观事实和数据"
        },
        "explanation": {
            "name": "解释",
            "color": "green",
            "keywords": ["因为", "所以", "原因", "由于", "导致", "说明", "解释", "原理", "机制", "逻辑"],
            "description": "原因分析和解释"
        },
        "risk": {
            "name": "风险",
            "color": "yellow",
            "keywords": ["风险", "问题", "挑战", "困难", "障碍", "限制", "不足", "缺陷", "隐患", "警告"],
            "description": "潜在风险和问题"
        },
        "action": {
            "name": "行动",
            "color": "red",
            "keywords": ["建议", "措施", "方案", "计划", "行动", "实施", "执行", "推进", "落实", "需要"],
            "description": "行动建议和对策"
        }
    }
    
    def __init__(self):
        """初始化处理器"""
        if not PDF_AVAILABLE:
            logger.warning("PDF 库未安装，PDF 功能不可用")
        elif not PYPDF_AVAILABLE:
            logger.warning("pypdf 未安装，PDF 功能受限")
        if not PDFPLUMBER_AVAILABLE:
            logger.warning("pdfplumber 未安装，表格提取功能受限")
    
    def generate_four_color_cards(self, pdf_path: str, max_cards: int = 50) -> Dict[str, Any]:
        """
        从 PDF 生成四色卡片
        
        Args:
            pdf_path: PDF 文件路径
            max_cards: 最大卡片数量
            
        Returns:
            包含卡片列表和统计信息的字典
        """
        result = {
            "success": False,
            "cards": [],
            "stats": {},
            "metadata": {},
            "error": None,
            "message": ""
        }
        
        try:
            # 提取 PDF 文本
            text_content = self._extract_text(pdf_path)
            
            if not text_content:
                result["error"] = "无法从 PDF 提取文本"
                return result
            
            # 分析内容并生成卡片
            cards = self._analyze_content(text_content, max_cards)
            
            # 统计信息
            stats = {
                "total_cards": len(cards),
                "fact_cards": len([c for c in cards if c["type"] == "fact"]),
                "explanation_cards": len([c for c in cards if c["type"] == "explanation"]),
                "risk_cards": len([c for c in cards if c["type"] == "risk"]),
                "action_cards": len([c for c in cards if c["type"] == "action"])
            }
            
            # 元数据
            metadata = {
                "source_file": os.path.basename(pdf_path),
                "total_length": len(text_content),
                "extraction_method": "pypdf" if PYPDF_AVAILABLE else "fallback"
            }
            
            result["success"] = True
            result["cards"] = cards
            result["stats"] = stats
            result["metadata"] = metadata
            result["message"] = f"成功生成 {len(cards)} 张四色卡片"
            
            logger.info(f"PDF 四色卡片生成完成: {len(cards)} 张卡片")
            
        except Exception as e:
            result["error"] = str(e)
            logger.error(f"PDF 四色卡片生成失败: {e}", exc_info=True)
        
        return result
    
    def _extract_text(self, pdf_path: str) -> str:
        """
        从 PDF 提取文本
        
        Args:
            pdf_path: PDF 文件路径
            
        Returns:
            提取的文本内容
        """
        text_content = []
        
        try:
            if PYPDF_AVAILABLE:
                reader = PdfReader(pdf_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            elif PDFPLUMBER_AVAILABLE:
                # 降级方案：尝试使用 pdfplumber
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
            else:
                logger.error("没有可用的 PDF 库，无法提取文本")
                return ""
        except Exception as e:
            logger.error(f"PDF 文本提取失败: {e}")
        
        return "\n".join(text_content)
    
    def _analyze_content(self, text: str, max_cards: int) -> List[Dict[str, Any]]:
        """
        分析内容并生成四色卡片
        
        Args:
            text: 文本内容
            max_cards: 最大卡片数量
            
        Returns:
            卡片列表
        """
        cards = []
        
        # 按段落分割文本
        paragraphs = self._split_into_paragraphs(text)
        
        # 分析每个段落
        for i, paragraph in enumerate(paragraphs):
            if len(cards) >= max_cards:
                break
            
            if len(paragraph) < 20:  # 跳过太短的段落
                continue
            
            # 判断段落类型
            card_type = self._classify_paragraph(paragraph)
            
            # 生成标题
            title = self._generate_title(paragraph, card_type)
            
            # 创建卡片
            card = {
                "id": f"card-{i+1}",
                "type": card_type,
                "title": title,
                "content": paragraph[:500] + "..." if len(paragraph) > 500 else paragraph,
                "tags": self._extract_tags(paragraph),
                "source": "PDF分析",
                "confidence": 0.8
            }
            
            cards.append(card)
        
        return cards
    
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """
        将文本分割成段落
        
        Args:
            text: 文本内容
            
        Returns:
            段落列表
        """
        # 按换行符分割
        paragraphs = re.split(r'\n\s*\n', text)
        
        # 清理每个段落
        cleaned_paragraphs = []
        for p in paragraphs:
            p = p.strip()
            if len(p) > 30:  # 只保留有意义的段落
                cleaned_paragraphs.append(p)
        
        return cleaned_paragraphs
    
    def _classify_paragraph(self, paragraph: str) -> str:
        """
        分类段落类型
        
        Args:
            paragraph: 段落文本
            
        Returns:
            卡片类型
        """
        scores = {card_type: 0 for card_type in self.CARD_TYPES.keys()}
        
        for card_type, config in self.CARD_TYPES.items():
            for keyword in config["keywords"]:
                if keyword in paragraph:
                    scores[card_type] += 1
        
        # 返回得分最高的类型
        max_type = max(scores, key=scores.get)
        
        # 如果没有明显特征，默认为 fact
        if scores[max_type] == 0:
            # 检查是否包含数字
            if re.search(r'\d+', paragraph):
                return "fact"
            return "explanation"
        
        return max_type
    
    def _generate_title(self, paragraph: str, card_type: str) -> str:
        """
        生成卡片标题
        
        Args:
            paragraph: 段落文本
            card_type: 卡片类型
            
        Returns:
            标题
        """
        # 取前30个字符作为标题
        title = paragraph[:30].strip()
        
        # 如果标题以标点符号结尾，去掉它
        title = re.sub(r'[，。！？；：""''（）【】]$', '', title)
        
        # 添加类型前缀
        type_name = self.CARD_TYPES[card_type]["name"]
        
        if len(title) < 10:
            return f"{type_name}: {title}"
        
        return title
    
    def _extract_tags(self, paragraph: str) -> List[str]:
        """
        提取标签
        
        Args:
            paragraph: 段落文本
            
        Returns:
            标签列表
        """
        tags = []
        
        # 简单的关键词提取
        common_tags = ["重要", "关键", "核心", "主要", "基础", "高级", "初级"]
        for tag in common_tags:
            if tag in paragraph:
                tags.append(tag)
        
        return tags[:3]  # 最多3个标签
    
    def _sanitize_for_docx(self, text: str) -> str:
        """移除 XML 1.0 非法字符，保留 tab/newline/cr"""
        if not isinstance(text, str):
            text = str(text)
        # XML 1.0 允许: 0x09 0x0A 0x0D, 0x20-0xD7FF, 0xE000-0xFFFD, 0x10000-0x10FFFF
        def _valid(c: str) -> bool:
            cp = ord(c)
            if cp in (0x09, 0x0A, 0x0D):
                return True
            if 0x20 <= cp <= 0xD7FF:
                return True
            if 0xE000 <= cp <= 0xFFFD:
                return True
            if 0x10000 <= cp <= 0x10FFFF:
                return True
            return False
        return ''.join(c if _valid(c) else ' ' for c in text)
    
    def export_to_excel(self, cards: List[Dict[str, Any]], output_path: str,
                       title: str = "四色知识卡片报告") -> Dict[str, Any]:
        """
        将卡片导出为 Excel（集成 skills/xlsx 专业导出器）
        
        支持多工作表：概览 + 按类型分组的卡片详情
        """
        result = {"success": False, "output_path": output_path, "error": None}

        if not EXCEL_AVAILABLE:
            result["error"] = "openpyxl 未安装，请运行: pip install openpyxl"
            return result

        try:
            # 尝试使用 skills/xlsx 的专业导出器（多工作表、概览页、图表支持）
            try:
                from skills.xlsx.excel_exporter import AntinetExcelExporter

                exporter = AntinetExcelExporter()
                exporter.create_workbook()

                # 统计卡片数量
                cards_by_type: Dict[str, List[Dict]] = {}
                card_counts = {}
                for card in cards:
                    # 跳过非字典项（防御性过滤）
                    if not isinstance(card, dict):
                        logger.warning(f"[export_to_excel] 跳过非字典卡片: {type(card).__name__}")
                        continue
                    ct = card.get("type", "fact")
                    cards_by_type.setdefault(ct, []).append(card)
                    type_name = exporter.CARD_NAMES.get(ct, ct)
                    card_counts[type_name] = card_counts.get(type_name, 0) + 1

                # 概览工作表
                from datetime import datetime as dt
                # 净化卡片文本（移除 XML 非法控制字符）
                sanitized_cards_by_type: Dict[str, List[Dict]] = {}
                for ct, gcards in cards_by_type.items():
                    sanitized_cards_by_type[ct] = [
                        {k: self._sanitize_for_docx(v) if isinstance(v, str) else v for k, v in card.items()}
                        for card in gcards
                    ]

                analysis_info = {
                    "title": self._sanitize_for_docx(title),
                    "date": dt.now().strftime('%Y-%m-%d %H:%M'),
                    "data_source": "PDF 四色卡片分析",
                    "card_counts": card_counts,
                    "summary": f"共提取 {len(cards)} 张知识卡片，涵盖 {len(cards_by_type)} 种类型。",
                }
                exporter.add_overview_sheet(analysis_info)

                # 按类型分组的卡片工作表
                for ct, grouped_cards in sanitized_cards_by_type.items():
                    if grouped_cards:
                        exporter.add_cards_sheet(ct, grouped_cards)

                exporter.wb.save(output_path)
                result["success"] = True
                logger.info(f"Excel 导出成功（专业版 - 多工作表）: {output_path}")
                return result

            except ImportError:
                # skills/xlsx 不可用，回退到内置实现
                pass

            # ========== 内置回退实现（单工作表，但增强样式）==========
            from openpyxl import Workbook
            from openpyxl.styles import Border, Side

            wb = Workbook()
            ws = wb.active
            ws.title = "四色卡片"

            # 定义颜色方案（与 minimax-xlsx 技能一致）
            color_map = {
                "fact":       ("🔵 事实",   "D6DCE4", "0052CC"),
                "explanation":("🟢 解释",   "D9EAD3", "009900"),
                "risk":       ("🟡 风险",   "FFF2CC", "FFC000"),
                "action":     ("🔴 行动",   "FCE4D6", "C00000"),
                "blue":       ("🔵 核心概念","D6DCE4", "0052CC"),
                "green":      ("🟢 关联链接","D9EAD3", "009900"),
                "yellow":     ("🟡 参考来源","FFF2CC", "FFC000"),
                "red":        ("🔴 索引关键词","FCE4D6", "C00000"),
            }

            thin_border = Border(
                left=Side(style='thin', color='CCCCCC'),
                right=Side(style='thin', color='CCCCCC'),
                top=Side(style='thin', color='CCCCCC'),
                bottom=Side(style='thin', color='CCCCCC'),
            )

            # 表头
            headers = ["序号", "类型", "标题", "内容摘要", "标签", "来源"]
            for col, h in enumerate(headers, 1):
                cell = ws.cell(1, col, h)
                cell.font = Font(bold=True, color="FFFFFF", size=10)
                cell.fill = PatternFill(start_color="4472C4", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = thin_border

            # 数据行
            for idx, card in enumerate(cards, 1):
                ct = card.get("type", "fact")
                type_name, bg_color, font_color = color_map.get(ct, ("未知", "FFFFFF", "333333"))

                row_data = [
                    idx,
                    type_name,
                    self._sanitize_for_docx(card.get("title", "")),
                    self._sanitize_for_docx((card.get("content", "") or "")[:200] + ("..." if len(card.get("content", "")) > 200 else "")),
                    ", ".join(self._sanitize_for_docx(str(t)) for t in card.get("tags", [])),
                    self._sanitize_for_docx(card.get("source", "") or card.get("address", ""))
                ]
                ws.append(row_data)

                row_num = ws.max_row
                for col_idx in range(1, len(row_data) + 1):
                    cell = ws.cell(row_num, col_idx)
                    cell.fill = PatternFill(start_color=bg_color, fill_type="solid")
                    cell.border = thin_border
                    if col_idx == 3:
                        cell.font = Font(bold=True, size=10)
                    else:
                        cell.font = Font(size=9.5)
                    cell.alignment = Alignment(vertical="top", wrap_text=True)

            # 列宽
            widths = [8, 14, 30, 50, 18, 20]
            for i, w in enumerate(widths, 1):
                ws.column_dimensions[chr(64 + i)].width = w

            # 自动筛选
            ws.auto_filter.ref = f"A1:F{ws.max_row}"

            # 冻结首行
            ws.freeze_panes = 'A2'

            wb.save(output_path)
            result["success"] = True
            logger.info(f"Excel 导出成功（标准版）: {output_path}, 共 {len(cards)} 张卡片")

        except Exception as e:
            result["error"] = str(e)
            logger.error(f"Excel 导出失败: {e}", exc_info=True)

        return result

    def export_to_docx(self, cards: List[Dict[str, Any]], output_path: str,
                      title: str = "四色知识卡片报告", author: str = "Antinet 智能知识管家") -> Dict[str, Any]:
        """将卡片导出为 Word 文档（专业排版，参考 minimax-docx 技能规范）"""
        result = {"success": False, "output_path": output_path, "error": None}

        if not DOCX_AVAILABLE:
            result["error"] = "python-docx 未安装，请运行: pip install python-docx"
            return result
        
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document()

            # ========== 文档属性（参考 GB/T 9704-2012 公文标准）==========
            core_props = doc.core_properties
            core_props.title = title
            core_props.author = author

            # ========== 页面设置 ==========
            section = doc.sections[0]
            section.page_width = Inches(8.27)    # A4
            section.page_height = Inches(11.69)  # A4
            section.left_margin = Inches(0.79)   # 2cm
            section.right_margin = Inches(0.79)
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.59)

            # ========== 页眉页脚 ==========
            title = self._sanitize_for_docx(title)
            author = self._sanitize_for_docx(author)
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{title}"
            header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
            if not header_para.runs:
                header_para.add_run(f"{title}")
            header_para.runs[0].font.size = Pt(9)
            header_para.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"第 "
            run1 = footer_para.add_run()
            # 使用域代码实现页码
            def add_page_number(paragraph):
                fld_char_begin = OxmlElement('w:fldChar')
                fld_char_begin.set(qn('w:fldCharType'), 'begin')
                instr_text = OxmlElement('w:instrText')
                instr_text.text = "PAGE"
                fld_char_end = OxmlElement('w:fldChar')
                fld_char_end.set(qn('w:fldCharType'), 'end')
                run = paragraph.add_run()
                run._r.append(fld_char_begin)
                run._r.append(instr_text)
                run._r.append(fld_char_end)
                return run
            
            # 清空默认页脚内容，重新构建
            footer_para.clear()
            pn_run = add_page_number(footer_para)
            pn_run.font.size = Pt(9)
            pn_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            sep_run = footer_para.add_run(" / ")
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            tp_run = add_page_number(footer_para)
            tp_run.font.size = Pt(9)
            tp_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            author_run = footer_para.add_run(f"  |  {author}  |  Antinet 智能知识管家")
            author_run.font.size = Pt(9)
            author_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ========== 颜色方案（参考四色卡片标准）==========
            color_map = {
                "fact":       ("事实",     RGBColor(0x00, 0x52, 0xCC), "DDEBF7"),
                "explanation":("解释",     RGBColor(0x00, 0x99, 0x00), "E2EFDA"),
                "risk":       ("风险",     RGBColor(0xFF, 0xC0, 0x00), "FFF2CC"),
                "action":     ("行动",     RGBColor(0xC0, 0x00, 0x00), "FCE4D6"),
                # 兼容英文类型名
                "blue":       ("核心概念", RGBColor(0x00, 0x52, 0xCC), "DDEBF7"),
                "green":      ("关联链接", RGBColor(0x00, 0x99, 0x00), "E2EFDA"),
                "yellow":     ("参考来源", RGBColor(0xFF, 0xC0, 0x00), "FFF2CC"),
                "red":        ("索引关键词",RGBColor(0xC0, 0x00, 0x00), "FCE4D6"),
            }

            # ========== 封面标题 ==========
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = subtitle.add_run(f"{author}  |  共 {len(cards)} 张知识卡片")
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # 分隔线效果：添加一个段落作为视觉分隔
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)

            # ========== 统计摘要表格 ==========
            type_count: Dict[str, int] = {}
            cards_by_type: Dict[str, List[Dict]] = {}
            for card in cards:
                ct = card.get("type", "fact")
                type_name, _, _ = color_map.get(ct, ("其他", RGBColor(0, 0, 0), "FFFFFF"))
                type_count[type_name] = type_count.get(type_name, 0) + 1
                cards_by_type.setdefault(ct, []).append(card)

            if len(cards) > 0 and type_count:
                stats_heading = doc.add_paragraph()
                sh_run = stats_heading.add_run("📊 卡片统计")
                sh_run.font.size = Pt(13)
                sh_run.font.bold = True
                sh_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                
                stats_table = doc.add_table(rows=1, cols=4)
                stats_table.style = 'Table Grid'
                hdr_cells = stats_table.rows[0].cells
                display_order = ["事实", "解释", "风险", "行动"]
                col_idx = 0
                for tn in display_order:
                    count = type_count.get(tn, 0)
                    bg_color = None
                    for k, (name, rgb, hex_c) in color_map.items():
                        if name == tn:
                            bg_color = hex_c
                            break
                    cell = hdr_cells[col_idx]
                    cell.text = f"{tn}\n{count}"
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                    # 背景色
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), bg_color or "F0F0F0")
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    col_idx += 1

                doc.add_paragraph()  # 空行分隔

            # ========== 卡片内容（按类型分组展示）==========
            display_order_type = ["fact", "blue", "explanation", "green", "risk", "yellow", "action", "red"]

            for group_idx, ct in enumerate(display_order_type):
                grouped_cards = cards_by_type.get(ct, [])
                if not grouped_cards:
                    continue
                
                type_name, type_color, bg_hex = color_map.get(ct, ("其他", RGBColor(0, 0, 0), "FFFFFF"))

                # 类型标题
                type_heading = doc.add_paragraph()
                th_run = type_heading.add_run(f"■ {type_name} ({len(grouped_cards)}张)")
                th_run.font.size = Pt(14)
                th_run.font.bold = True
                th_run.font.color.rgb = type_color
                type_heading.paragraph_format.space_before = Pt(16)
                type_heading.paragraph_format.space_after = Pt(8)

                # 每张卡片
                for card in grouped_cards:
                    # 卡片标题
                    card_title = doc.add_paragraph()
                    ct_run = card_title.add_run(f"[{type_name}] ")
                    ct_run.font.bold = True
                    ct_run.font.color.rgb = type_color
                    ct_run.font.size = Pt(12)
                    
                    title_text = self._sanitize_for_docx(card.get("title", ""))
                    t_run = card_title.add_run(title_text)
                    t_run.font.bold = True
                    t_run.font.size = Pt(12)
                    t_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                    card_title.paragraph_format.space_before = Pt(10)
                    card_title.paragraph_format.space_after = Pt(3)

                    # 卡片内容
                    raw_content = self._sanitize_for_docx(card.get("content", ""))
                    content_text = raw_content[:1000] + ("..." if len(raw_content) > 1000 else "")

                    content_para = doc.add_paragraph()
                    crun = content_para.add_run(content_text)
                    crun.font.size = Pt(10.5)
                    crun.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    content_para.paragraph_format.space_after = Pt(2)

                    # 来源信息
                    source_text = ""
                    if card.get("source"):
                        source_text = f"来源：{self._sanitize_for_docx(card.get('source'))}"
                    elif card.get("address"):
                        source_text = f"来源：{self._sanitize_for_docx(card.get('address'))}"
                    elif card.get("tags"):
                        source_text = f"标签：{', '.join(self._sanitize_for_docx(str(t)) for t in card.get('tags', []))}"

                    if source_text:
                        meta_para = doc.add_paragraph()
                        mrun = meta_para.add_run(source_text)
                        mrun.font.size = Pt(8.5)
                        mrun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                        meta_para.paragraph_format.space_after = Pt(6)

            # ========== 文档末尾 ==========
            doc.add_paragraph()
            end_line = doc.add_paragraph()
            end_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            erun = end_line.add_run("— 由 Antinet 智能知识管家自动生成 —")
            erun.font.size = Pt(9)
            erun.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            erun.font.italic = True

            doc.save(output_path)
            result["success"] = True
            logger.info(f"Word 导出成功（专业排版）: {output_path}, 共 {len(cards)} 张卡片")
        except Exception as e:
            result["error"] = str(e)
            logger.error(f"Word 导出失败: {e}", exc_info=True)

        return result

    def export_to_pdf(self, cards: List[Dict[str, Any]], output_path: str,
                      title: str = "四色知识卡片报告", author: str = "Antinet 智能知识管家") -> Dict[str, Any]:
        """将卡片导出为 PDF 文档（纯 reportlab，无需系统依赖）"""
        result = {"success": False, "output_path": output_path, "error": None}

        if not REPORTLAB_AVAILABLE:
            result["error"] = "reportlab 未安装，请运行: pip install reportlab"
            return result

        try:
            font_name = _ensure_chinese_font()
            bold_font = f'{font_name}-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold'
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                leftMargin=2*cm, rightMargin=2*cm,
                topMargin=2*cm, bottomMargin=2*cm,
                title=title,
                author=author,
            )

            # 颜色定义 (r, g, b) — 与四色卡片一致
            COLOR_BLUE   = colors.HexColor('#0052CC')  # 核心概念/事实
            COLOR_GREEN  = colors.HexColor('#009900')  # 关联链接/解释
            COLOR_YELLOW = colors.HexColor('#FFC000')  # 参考来源/风险
            COLOR_RED    = colors.HexColor('#C00000')  # 索引关键词/行动

            type_color_map = {
                "fact":       ("事实",       COLOR_BLUE),
                "explanation":("解释",        COLOR_GREEN),
                "risk":       ("风险",        COLOR_YELLOW),
                "action":     ("行动",        COLOR_RED),
                # 兼容中文类型
                "blue":       ("核心概念",   COLOR_BLUE),
                "green":      ("关联链接",   COLOR_GREEN),
                "yellow":     ("参考来源",   COLOR_YELLOW),
                "red":        ("索引关键词", COLOR_RED),
            }

            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='DocTitle',
                parent=styles['Title'],
                fontSize=22,
                leading=28,
                fontName=bold_font,
                textColor=colors.HexColor('#1a1a2e'),
                alignment=TA_CENTER,
                spaceAfter=6,
            ))
            styles.add(ParagraphStyle(
                name='DocSubtitle',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                fontName=font_name,
                textColor=colors.HexColor('#666666'),
                alignment=TA_CENTER,
                spaceAfter=20,
            ))
            styles.add(ParagraphStyle(
                name='CardTitle',
                parent=styles['Normal'],
                fontSize=13,
                leading=18,
                fontName=bold_font,
                textColor=colors.HexColor('#1a1a2e'),
                spaceAfter=4,
            ))
            styles.add(ParagraphStyle(
                name='CardContent',
                parent=styles['Normal'],
                fontSize=10,
                leading=15,
                fontName=font_name,
                textColor=colors.HexColor('#333333'),
                spaceAfter=4,
            ))
            styles.add(ParagraphStyle(
                name='CardMeta',
                parent=styles['Normal'],
                fontSize=8,
                leading=11,
                fontName=font_name,
                textColor=colors.HexColor('#999999'),
            ))

            story = []

            # 封面标题
            story.append(Paragraph(title, styles['DocTitle']))
            story.append(Paragraph(f"作者：{author}  |  共 {len(cards)} 张知识卡片", styles['DocSubtitle']))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#e0e0e0'), spaceAfter=16))

            # 统计摘要
            type_count: Dict[str, int] = {}
            for card in cards:
                ct = card.get("type", "fact")
                type_name, _ = type_color_map.get(ct, ("其他", COLOR_BLUE))
                type_count[type_name] = type_count.get(type_name, 0) + 1

            if type_count:
                stats_row = [
                    Paragraph(f"<b><font color='#0052CC'>核心概念 {type_count.get('核心概念', 0)}</font></b>", styles['CardMeta']),
                    Paragraph(f"<b><font color='#009900'>关联链接 {type_count.get('关联链接', 0)}</font></b>", styles['CardMeta']),
                    Paragraph(f"<b><font color='#FFC000'>参考来源 {type_count.get('参考来源', 0)}</font></b>", styles['CardMeta']),
                    Paragraph(f"<b><font color='#C00000'>索引关键词 {type_count.get('索引关键词', 0)}</font></b>", styles['CardMeta']),
                ]
                stats_table = Table([stats_row], colWidths=[4*cm, 4*cm, 4*cm, 4*cm])
                stats_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f8f8')),
                    ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#f8f8f8'), colors.HexColor('#f0f0f0')]),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    ('ROUNDEDCORNERS', [4, 4, 4, 4]),
                ]))
                story.append(stats_table)
                story.append(Spacer(1, 16))

            # 卡片列表
            for i, card in enumerate(cards):
                ct = card.get("type", "fact")
                type_name, type_color = type_color_map.get(ct, ("其他", COLOR_BLUE))

                # 颜色标识条
                color_bar = Table([['']], colWidths=[18], rowHeights=[52])
                color_bar.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), type_color),
                    ('VALIGN', (0, 0), (0, 0), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (0, 0), 0),
                    ('BOTTOMPADDING', (0, 0), (0, 0), 0),
                    ('LEFTPADDING', (0, 0), (0, 0), 0),
                    ('RIGHTPADDING', (0, 0), (0, 0), 0),
                ]))

                content_text = card.get("content", "")
                # 截断过长内容
                if len(content_text) > 800:
                    content_text = content_text[:800] + "..."

                source_text = ""
                if card.get("source"):
                    source_text = f"来源：{card.get('source')}"
                elif card.get("address"):
                    source_text = f"来源：{card.get('address')}"

                # 标题行
                title_para = Paragraph(
                    f'<font color="#{type_color.hexval()[2:]}"><b>[{type_name}]</b></font>  {card.get("title", "")}',
                    styles['CardTitle']
                )
                content_para = Paragraph(content_text.replace('\n', '<br/>'), styles['CardContent'])

                card_content = [title_para, content_para]
                if source_text:
                    card_content.append(Paragraph(source_text, styles['CardMeta']))

                # 卡片行：[颜色条 | 内容]
                card_table = Table(
                    [[color_bar, card_content]],
                    colWidths=[0.5*cm, 16.5*cm],
                )
                card_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (0, 0), 0),
                    ('RIGHTPADDING', (0, 0), (0, 0), 0),
                    ('LEFTPADDING', (1, 0), (1, 0), 10),
                    ('RIGHTPADDING', (1, 0), (1, 0), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))

                story.append(card_table)
                story.append(Spacer(1, 8))

                # 分页
                if (i + 1) % 5 == 0 and i < len(cards) - 1:
                    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e0e0e0'), spaceAfter=8))
                    story.append(Spacer(1, 8))

            # 页脚
            story.append(Spacer(1, 20))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#cccccc'), spaceAfter=6))
            footer_style = ParagraphStyle(
                name='Footer',
                parent=styles['Normal'],
                fontSize=8,
                fontName=font_name,
                textColor=colors.HexColor('#aaaaaa'),
                alignment=TA_CENTER,
            )
            story.append(Paragraph("由 Antinet 智能知识管家自动生成  |  八府巡按，各司其职", footer_style))

            doc.build(story)
            result["success"] = True
            logger.info(f"PDF 导出成功: {output_path}，共 {len(cards)} 张卡片")

        except Exception as e:
            result["error"] = str(e)
            logger.error(f"PDF 导出失败: {e}", exc_info=True)

        return result
def generate_four_color_cards(pdf_path: str, max_cards: int = 50) -> Dict[str, Any]:
    """
    便捷函数：从 PDF 生成四色卡片
    
    Args:
        pdf_path: PDF 文件路径
        max_cards: 最大卡片数量
        
    Returns:
        包含卡片列表和统计信息的字典
    """
    processor = PDFourColorProcessor()
    return processor.generate_four_color_cards(pdf_path, max_cards)


def export_cards_to_excel(cards: List[Dict[str, Any]], output_path: str) -> Dict[str, Any]:
    """
    便捷函数：将卡片导出为 Excel
    
    Args:
        cards: 卡片列表
        output_path: 输出文件路径
        
    Returns:
        导出结果
    """
    processor = PDFourColorProcessor()
    return processor.export_to_excel(cards, output_path)


# 测试代码
if __name__ == "__main__":
    # 测试处理器
    processor = PDFourColorProcessor()
    
    # 测试文本分析
    test_text = """
    2024年公司营收达到1000万元，同比增长30%。
    
    这一增长主要得益于新产品的推出和市场扩张策略的实施。
    
    需要注意的是，市场竞争加剧可能带来一定的风险。
    
    建议加大研发投入，提升产品竞争力。
    """
    
    cards = processor._analyze_content(test_text, 10)
    
    print(f"生成 {len(cards)} 张卡片:")
    for card in cards:
        print(f"\n[{card['type']}] {card['title']}")
        print(f"  {card['content'][:100]}...")
