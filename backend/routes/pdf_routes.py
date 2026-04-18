"""
PDF 处理 API 路由
为 Antinet 提供 PDF 文档处理接口
"""

from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse
from typing import List, Optional
import os
import tempfile
from pathlib import Path
import shutil
import zipfile

# 延迟导入避免pypdf与pydantic冲突
from tools.pdf_processor import SimplePDFProcessor

# 使用稳定的四色卡片处理器来处理所有PDF操作
try:
    from tools.pdf_four_color_processor import PDFourColorProcessor
    FOUR_COLOR_AVAILABLE = True
except ImportError:
    FOUR_COLOR_AVAILABLE = False

router = APIRouter(prefix="/api/pdf", tags=["PDF处理"])

# 初始化 PDF 处理器
pdf_processor = SimplePDFProcessor()

# 初始化四色卡片处理器
if FOUR_COLOR_AVAILABLE and PDF_AVAILABLE:
    four_color_processor = PDFourColorProcessor()
else:
    four_color_processor = None


@router.get("/status")
async def get_pdf_status():
    """获取 PDF 功能状态（含 MinerU 高质量解析能力）"""
    base_status = {
        "available": PDF_AVAILABLE,
        "message": "PDF 功能已启用" if PDF_AVAILABLE else "PDF 功能未安装，请运行: pip install pypdf pdfplumber reportlab",
        "four_color_available": FOUR_COLOR_AVAILABLE and PDF_AVAILABLE,
        "toolkit_available": PDF_TOOLKIT_AVAILABLE,
    }

    # MinerU 状态
    try:
        from tools.mineru_processor import get_mineru_processor
        proc = get_mineru_processor()
        ms = proc.get_status()
        backends = []
        if ms["vlm_client_available"]:
            backends.append("vlm-http-client")
        if ms["pipeline_available"]:
            backends.append("pipeline")
        if ms["cli_path"]:
            backends.append("cli")
        base_status["mineru"] = {
            "available": ms["available"],
            "version": ms["version"],
            "backends": backends,
            "message": "MinerU 已就绪 - 支持高质量 PDF→Markdown+JSON" if ms["available"] else "MinerU 未安装",
        }
    except Exception as e:
        base_status["mineru"] = {
            "available": False,
            "version": None,
            "backends": [],
            "message": f"MinerU 不可用: {e}",
        }

    return base_status


@router.post("/extract/text")
async def extract_text(
    file: UploadFile = File(...),
    preserve_layout: bool = True
):
    """
    从 PDF 提取文本
    
    Args:
        file: PDF 文件
        preserve_layout: 是否保留布局
        
    Returns:
        提取的文本内容和元数据
    """
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 功能未安装")
    
    # 保存上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name
    
    try:
        # 提取文本
        result = pdf_processor.extract_text(tmp_path, preserve_layout)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return {
            "success": True,
            "filename": file.filename,
            "pages": len(result["pages"]),
            "text": result["full_text"],
            "full_text": result["full_text"],
            "metadata": result["metadata"]
        }
    
    finally:
        # 清理临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.post("/extract/tables")
async def extract_tables(
    file: UploadFile = File(...),
    page_numbers: Optional[str] = Form(None)
):
    """
    从 PDF 提取表格
    
    Args:
        file: PDF 文件
        page_numbers: 页码列表 (逗号分隔, 如 "1,2,3")
        
    Returns:
        提取的表格数据
    """
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 功能未安装")
    
    # 保存上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name
    
    try:
        # 解析页码
        pages = None
        if page_numbers:
            try:
                pages = [int(p.strip()) for p in page_numbers.split(",")]
            except ValueError:
                raise HTTPException(status_code=400, detail="页码格式错误")
        
        # 提取表格
        result = pdf_processor.extract_tables(tmp_path)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return {
            "success": True,
            "filename": file.filename,
            "tables": result["tables"]
        }
    
    finally:
        # 清理临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.post("/extract/knowledge")
async def extract_knowledge(file: UploadFile = File(...)):
    """
    从 PDF 提取知识并准备生成四色卡片
    
    Args:
        file: PDF 文件
        
    Returns:
        提取的知识内容和建议的卡片类型
    """
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 功能未安装")
    
    # 保存上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name
    
    try:
        # 提取知识
        result = pdf_processor.extract_knowledge(tmp_path)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return {
            "success": True,
            "filename": file.filename,
            "text_content": result["text_content"],
            "tables": result["tables"],
            "metadata": result["metadata"],
            "suggested_cards": result["suggested_cards"]
        }
    
    finally:
        # 清理临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.post("/export/cards")
async def export_cards(
    cards: List[dict],
    title: str = "Antinet 分析报告",
    author: str = "Antinet 智能知识管家"
):
    """
    将四色卡片导出为 PDF 报告
    
    Args:
        cards: 四色卡片列表
        title: 报告标题
        author: 报告作者
        
    Returns:
        生成的 PDF 文件
    """
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 功能未安装")
    
    if not cards:
        raise HTTPException(status_code=400, detail="卡片列表不能为空")
    
    # 创建临时输出文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        output_path = tmp_file.name
    
    try:
        # 导出 PDF
        result = pdf_processor.export_cards_to_pdf(cards, output_path, title, author)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        # 返回文件
        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename=f"{title}_{result['cards_count']}cards.pdf"
        )
    
    except Exception as e:
        # 清理临时文件
        if os.path.exists(output_path):
            os.unlink(output_path)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/batch/process")
async def batch_process(
    files: List[UploadFile] = File(...),
    extract_text: bool = True,
    extract_tables: bool = True
):
    """
    批量处理 PDF 文档
    
    Args:
        files: PDF 文件列表
        extract_text: 是否提取文本
        extract_tables: 是否提取表格
        
    Returns:
        批量处理结果
    """
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 功能未安装")
    
    if not files:
        raise HTTPException(status_code=400, detail="文件列表不能为空")
    
    # 创建临时目录
    with tempfile.TemporaryDirectory() as tmp_dir:
        input_dir = os.path.join(tmp_dir, "input")
        output_dir = os.path.join(tmp_dir, "output")
        os.makedirs(input_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)
        
        # 保存上传的文件
        for file in files:
            file_path = os.path.join(input_dir, file.filename)
            with open(file_path, "wb") as f:
                shutil.copyfileobj(file.file, f)
        
        # 批量处理（简化版本）
        processed_results = []
        failed_results = []
        
        for filename in os.listdir(input_dir):
            file_path = os.path.join(input_dir, filename)
            try:
                if extract_text:
                    result = pdf_processor.extract_text(file_path, True)
                    if result.get("success"):
                        processed_results.append({
                            "filename": filename,
                            "status": "success",
                            "text_extracted": True
                        })
                    else:
                        failed_results.append({
                            "filename": filename,
                            "status": "failed",
                            "error": result.get("error", "无法提取文本")
                        })
                else:
                    processed_results.append({
                        "filename": filename,
                        "status": "success",
                        "text_extracted": False
                    })
            except Exception as e:
                failed_results.append({
                    "filename": filename,
                    "status": "failed",
                    "error": str(e)
                })
        
        return {
            "success": True,
            "total": len(files),
            "processed": len(processed_results),
            "failed": len(failed_results),
            "results": processed_results + failed_results
        }


# ========== 四色卡片生成 ==========

@router.post("/generate/four-color-cards")
async def generate_four_color_cards(
    file: UploadFile = File(...),
    max_cards: int = 50
):
    """
    从 PDF 生成四色卡片（智能分类）
    
    Args:
        file: PDF 文件
        max_cards: 最大卡片数量
        
    Returns:
        四色卡片列表和统计信息
    """
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 功能未安装")
    
    if not FOUR_COLOR_AVAILABLE or four_color_processor is None:
        raise HTTPException(status_code=503, detail="四色卡片处理器未安装")
    
    # 保存上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name
    
    try:
        # 生成四色卡片
        result = four_color_processor.generate_four_color_cards(tmp_path, max_cards)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return {
            "success": True,
            "filename": file.filename,
            "cards": result["cards"],
            "stats": result["stats"],
            "metadata": result["metadata"],
            "message": result["message"]
        }
    
    finally:
        # 清理临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.post("/export/four-color-excel")
async def export_four_color_excel(
    file: UploadFile = File(...),
    max_cards: int = 50
):
    """
    从 PDF 生成四色卡片并导出为 Excel
    
    Args:
        file: PDF 文件
        max_cards: 最大卡片数量
        
    Returns:
        四色卡片 Excel 文件
    """
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 功能未安装")
    
    if not FOUR_COLOR_AVAILABLE or four_color_processor is None:
        raise HTTPException(status_code=503, detail="四色卡片处理器未安装")
    
    # 保存上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name
    
    # 创建输出文件路径
    output_path = tmp_path.replace(".pdf", "_four_color_cards.xlsx")
    
    try:
        # 生成四色卡片
        cards_result = four_color_processor.generate_four_color_cards(tmp_path, max_cards)
        
        if not cards_result["success"]:
            raise HTTPException(status_code=500, detail=cards_result["error"])
        
        # 导出到 Excel
        export_result = four_color_processor.export_to_excel(cards_result["cards"], output_path)
        
        if not export_result["success"]:
            raise HTTPException(status_code=500, detail=export_result["error"])
        
        # 返回文件
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=f"{Path(file.filename).stem}_四色卡片.xlsx",
            background=None  # 确保文件发送后再清理
        )
    
    except Exception as e:
        # 清理临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        if os.path.exists(output_path):
            os.unlink(output_path)
        raise HTTPException(status_code=500, detail=str(e))


# ========== 健康检查 ==========

@router.get("/health")
async def health_check():
    """健康检查"""
    return {
        "status": "healthy" if PDF_AVAILABLE else "unavailable",
        "service": "PDF Processing",
        "version": "1.0.0"
    }


try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


from pydantic import BaseModel
from typing import Dict, Any


class CardsExportRequest(BaseModel):
    """卡片导出请求"""
    cards: List[Dict[str, Any]]
    title: str = "Antinet 分析报告"
    author: str = "Antinet 智能知识管家"


@router.post("/export/cards-docx")
async def export_cards_to_docx(request: CardsExportRequest):
    """
    将四色卡片导出为 Word 文档
    
    请求体:
    {
        "cards": [...],
        "title": "报告标题",
        "author": "作者"
    }
    """
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=503, detail="Word 导出功能未安装，请运行: pip install python-docx")
    
    if not request.cards:
        raise HTTPException(status_code=400, detail="卡片列表不能为空")
    
    try:
        doc = Document()
        
        doc.add_heading(request.title, 0)
        doc.add_paragraph(f"作者: {request.author}")
        doc.add_paragraph(f"生成时间: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        card_colors = {
            "blue": RGBColor(52, 152, 219),
            "green": RGBColor(46, 204, 113),
            "yellow": RGBColor(241, 196, 15),
            "red": RGBColor(231, 76, 60)
        }
        
        card_names = {
            "blue": "事实卡片",
            "green": "解释卡片",
            "yellow": "风险卡片",
            "red": "行动卡片"
        }
        
        for idx, card in enumerate(request.cards, 1):
            card_type = card.get("card_type", card.get("type", "blue"))
            card_title = card.get("title", f"卡片 {idx}")
            card_content = card.get("content", "")
            
            if isinstance(card_content, dict):
                card_content = card_content.get("description", str(card_content))
            
            heading = doc.add_heading(f"{idx}. {card_title}", level=1)
            
            color = card_colors.get(card_type, RGBColor(128, 128, 128))
            for run in heading.runs:
                run.font.color.rgb = color
            
            type_para = doc.add_paragraph()
            type_run = type_para.add_run(f"[{card_names.get(card_type, '卡片')}]")
            type_run.font.color.rgb = color
            type_run.font.bold = True
            
            doc.add_paragraph(card_content)
            doc.add_paragraph()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            output_path = tmp_file.name
        
        doc.save(output_path)
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{request.title}.docx"
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


# ==================== PDF 工具集路由 ====================

try:
    from tools.pdf_toolkit import PDFToolkit
    PDF_TOOLKIT_AVAILABLE = True
except ImportError:
    PDF_TOOLKIT_AVAILABLE = False


@router.post("/toolkit/merge")
async def merge_pdfs(
    files: List[UploadFile] = File(..., description="要合并的 PDF 文件列表")
):
    """
    合并多个 PDF 文件
    
    Args:
        files: PDF 文件列表（至少2个）
        
    Returns:
        合并后的 PDF 文件
    """
    if not PDF_TOOLKIT_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 工具集未安装")
    
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="至少需要2个 PDF 文件")
    
    # 保存上传的文件
    input_files = []
    for file in files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            shutil.copyfileobj(file.file, tmp_file)
            input_files.append(tmp_file.name)
    
    # 创建输出文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        output_path = tmp_file.name
    
    try:
        # 合并 PDF
        result = PDFToolkit.merge_pdfs(input_files, output_path)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename="merged.pdf"
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"合并失败: {str(e)}")
    
    finally:
        # 清理临时文件
        for f in input_files:
            try:
                os.unlink(f)
            except:
                pass


@router.post("/toolkit/split")
async def split_pdf(
    file: UploadFile = File(..., description="要拆分的 PDF 文件"),
    page_range: Optional[str] = Form(None, description="页码范围（如 '1,3,5-7'），为空则拆分为单页")
):
    """
    拆分 PDF 文件
    
    Args:
        file: PDF 文件
        page_range: 页码范围（如 "1,3,5-7"），为空则拆分为单页
        
    Returns:
        拆分后的 PDF 文件（ZIP 压缩包）
    """
    if not PDF_TOOLKIT_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 工具集未安装")
    
    # 保存上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        input_path = tmp_file.name
    
    # 创建输出目录
    output_dir = tempfile.mkdtemp()
    
    try:
        # 拆分 PDF
        result = PDFToolkit.split_pdf(input_path, output_dir, page_range)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        # 创建 ZIP 文件
        zip_path = os.path.join(tempfile.gettempdir(), "split_pdfs.zip")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for output_file in result["output_files"]:
                zipf.write(output_file, os.path.basename(output_file))
        
        return FileResponse(
            zip_path,
            media_type="application/zip",
            filename="split_pdfs.zip"
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"拆分失败: {str(e)}")
    
    finally:
        # 清理临时文件
        try:
            os.unlink(input_path)
            shutil.rmtree(output_dir)
        except:
            pass


@router.post("/toolkit/pdf-to-images")
async def pdf_to_images(
    file: UploadFile = File(..., description="PDF 文件"),
    format: str = Form("jpg", description="输出格式（jpg/png）"),
    dpi: int = Form(150, description="图片分辨率"),
    pages: Optional[str] = Form(None, description="页码范围（如 '1-3'），为空则转换所有页")
):
    """
    将 PDF 转换为图片
    
    Args:
        file: PDF 文件
        format: 输出格式（jpg/png）
        dpi: 图片分辨率
        pages: 页码范围（如 "1-3"），为空则转换所有页
        
    Returns:
        图片文件（ZIP 压缩包）
    """
    if not PDF_TOOLKIT_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 工具集未安装")
    
    # 保存上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        input_path = tmp_file.name
    
    # 创建输出目录
    output_dir = tempfile.mkdtemp()
    
    try:
        # 转换 PDF 为图片
        result = PDFToolkit.pdf_to_images(input_path, output_dir, format, dpi, pages)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        # 创建 ZIP 文件
        zip_path = os.path.join(tempfile.gettempdir(), "pdf_images.zip")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for output_file in result["output_files"]:
                zipf.write(output_file, os.path.basename(output_file))
        
        return FileResponse(
            zip_path,
            media_type="application/zip",
            filename="pdf_images.zip"
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")
    
    finally:
        # 清理临时文件
        try:
            os.unlink(input_path)
            shutil.rmtree(output_dir)
        except:
            pass


@router.post("/toolkit/images-to-pdf")
async def images_to_pdf(
    files: List[UploadFile] = File(..., description="图片文件列表（jpg/png/bmp/tiff）")
):
    """
    将多张图片合并为 PDF
    
    Args:
        files: 图片文件列表
        
    Returns:
        合并后的 PDF 文件
    """
    if not PDF_TOOLKIT_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 工具集未安装")
    
    if len(files) < 1:
        raise HTTPException(status_code=400, detail="至少需要1个图片文件")
    
    # 保存上传的文件
    input_files = []
    for file in files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp_file:
            shutil.copyfileobj(file.file, tmp_file)
            input_files.append(tmp_file.name)
    
    # 创建输出文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        output_path = tmp_file.name
    
    try:
        # 合并图片为 PDF
        result = PDFToolkit.images_to_pdf(input_files, output_path)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename="images.pdf"
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"合并失败: {str(e)}")
    
    finally:
        # 清理临时文件
        for f in input_files:
            try:
                os.unlink(f)
            except:
                pass


# ==================== MinerU 高质量解析路由 ====================
# MinerU 2.5/3.x：复杂PDF一键变完美Markdown+结构化JSON
# 公式、表格、阅读顺序全保真

try:
    from tools.mineru_processor import get_mineru_processor, MINERU_AVAILABLE as _MINERU_AVAILABLE
    MINERU_ROUTE_AVAILABLE = True
except ImportError:
    _MINERU_AVAILABLE = False
    MINERU_ROUTE_AVAILABLE = False


@router.get("/mineru/status")
async def get_mineru_status():
    """
    获取 MinerU 可用状态

    返回 MinerU 版本、支持的 backend 以及所需依赖状态。
    """
    if not MINERU_ROUTE_AVAILABLE:
        return {
            "available": False,
            "message": "MinerU 未安装，请运行: pip install mineru",
            "version": None,
            "backends": [],
        }
    proc = get_mineru_processor()
    status = proc.get_status()
    backends = []
    if status["vlm_client_available"]:
        backends.append("vlm-http-client")
    if status["pipeline_available"]:
        backends.append("pipeline")
    if status["cli_path"]:
        backends.append("cli")
    return {
        "available": status["available"],
        "version": status["version"],
        "backends": backends,
        "vlm_client_available": status["vlm_client_available"],
        "pipeline_available": status["pipeline_available"],
        "cli_path": status["cli_path"],
        "message": "MinerU 已就绪" if status["available"] else f"MinerU 不可用: {status['import_error']}",
    }


@router.post("/mineru/parse")
async def mineru_parse_pdf(
    file: UploadFile = File(..., description="PDF 文件"),
    backend: str = Form("auto", description="解析后端: auto | vlm-http-client | pipeline | cli"),
    language: str = Form("ch", description="文档语言: ch | en"),
    formula_enable: bool = Form(True, description="是否识别公式"),
    table_enable: bool = Form(True, description="是否识别表格"),
    vlm_server_url: Optional[str] = Form(None, description="VLM HTTP server 地址（backend=vlm-http-client 时使用）"),
    return_content_list: bool = Form(True, description="是否返回结构化 content_list JSON"),
    return_middle_json: bool = Form(False, description="是否返回中间 middle_json（体积较大）"),
):
    """
    🚀 MinerU 高质量 PDF 解析

    基于 MinerU 3.x，将复杂 PDF 一键转换为：
    - **Markdown**：公式、表格、阅读顺序全保真
    - **content_list.json**：结构化 JSON（段落/表格/图片/公式分类）
    - **middle.json**（可选）：完整中间语义层

    适合：学术论文、技术手册、财报、含公式/表格的复杂文档。

    **backend 选项：**
    - `auto`：自动选择（vlm-http-client > pipeline > cli）
    - `vlm-http-client`：调用外部 VLM API，需填 vlm_server_url
    - `pipeline`：本地模型（需要 torch + 已下载模型）
    - `cli`：通过 mineru CLI 子进程执行
    """
    if not MINERU_ROUTE_AVAILABLE or not _MINERU_AVAILABLE:
        # 检查是否可以通过 CLI 降级
        proc = get_mineru_processor() if MINERU_ROUTE_AVAILABLE else None
        if proc is None or not proc.get_status().get("cli_path"):
            raise HTTPException(
                status_code=503,
                detail="MinerU 未安装，请运行: pip install mineru"
            )

    proc = get_mineru_processor()

    # 保存上传文件
    suffix = Path(file.filename).suffix or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name

    # 创建专属输出目录（不自动删除，结果可复用）
    output_dir = tempfile.mkdtemp(prefix="mineru_out_")

    try:
        result = proc.parse_pdf(
            pdf_path=tmp_path,
            output_dir=output_dir,
            backend=backend,
            language=language,
            formula_enable=formula_enable,
            table_enable=table_enable,
            vlm_server_url=vlm_server_url or None,
        )

        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])

        response = {
            "success": True,
            "filename": file.filename,
            "backend_used": result["backend_used"],
            "markdown": result["markdown"],
            "markdown_length": len(result["markdown"]),
        }

        if return_content_list:
            response["content_list"] = result["content_list"]
            response["content_list_v2"] = result["content_list_v2"]
            response["content_list_count"] = len(result["content_list"])

        if return_middle_json:
            response["middle_json"] = result["middle_json"]

        return response

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"MinerU 解析失败: {str(e)}")

    finally:
        # 清理上传的临时文件（保留输出目录供下载）
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.post("/mineru/parse-to-zip")
async def mineru_parse_pdf_to_zip(
    file: UploadFile = File(..., description="PDF 文件"),
    backend: str = Form("auto"),
    language: str = Form("ch"),
    formula_enable: bool = Form(True),
    table_enable: bool = Form(True),
    vlm_server_url: Optional[str] = Form(None),
):
    """
    🚀 MinerU 解析 PDF → 下载完整结果 ZIP 包

    ZIP 包含：
    - `<name>.md`：Markdown 全文
    - `<name>_content_list.json`：结构化 JSON
    - `<name>_content_list_v2.json`：结构化 JSON v2
    - `<name>_middle.json`：中间语义层
    - `images/`：文档中的图片
    """
    if not MINERU_ROUTE_AVAILABLE:
        raise HTTPException(status_code=503, detail="MinerU 未安装")

    proc = get_mineru_processor()

    suffix = Path(file.filename).suffix or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name

    output_dir = tempfile.mkdtemp(prefix="mineru_out_")

    try:
        result = proc.parse_pdf(
            pdf_path=tmp_path,
            output_dir=output_dir,
            backend=backend,
            language=language,
            formula_enable=formula_enable,
            table_enable=table_enable,
            vlm_server_url=vlm_server_url or None,
        )

        if not result["success"]:
            raise HTTPException(status_code=500, detail=result["error"])

        # 打包为 ZIP
        zip_path = os.path.join(tempfile.gettempdir(), f"mineru_{Path(file.filename).stem}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            out_base = Path(output_dir)
            for f_path in out_base.rglob("*"):
                if f_path.is_file():
                    arcname = f_path.relative_to(out_base)
                    zf.write(f_path, arcname)

        return FileResponse(
            zip_path,
            media_type="application/zip",
            filename=f"{Path(file.filename).stem}_mineru.zip",
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"MinerU 解析失败: {str(e)}")

    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        shutil.rmtree(output_dir, ignore_errors=True)