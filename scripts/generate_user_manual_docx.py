#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Antinet 知易智能知识管家 - 用户操作手册', 0)
doc.add_paragraph('版本: 1.0.0')
doc.add_paragraph('日期: 2026-02-17')
doc.add_paragraph()

doc.add_heading('1 系统概述', 1)
doc.add_heading('1.1 产品介绍', 2)
doc.add_paragraph('Antinet知易智能知识管家是一款基于骁龙X Elite AIPC平台的端侧智能数据中枢与协同分析平台。')

doc.add_heading('1.2 核心功能', 2)
features = ['四色卡片系统', '八大智能体', 'AI对话', '知识库', 'GTD任务', '文档处理', '视觉理解']
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2 快速开始', 1)
doc.add_heading('2.1 启动系统', 2)
doc.add_paragraph('方式一: 使用启动脚本')
doc.add_paragraph('方式二: 手动启动')
doc.add_paragraph('  cd backend && python main.py')
doc.add_paragraph('  cd src && npm run dev')

doc.add_heading('3 四色卡片使用', 1)
doc.add_heading('3.1 创建卡片', 2)
doc.add_paragraph('1. 点击"卡片"进入卡片管理页面')
doc.add_paragraph('2. 点击"新建卡片"')
doc.add_paragraph('3. 选择颜色类型（蓝/绿/黄/红）')
doc.add_paragraph('4. 填写标题和内容')
doc.add_paragraph('5. 点击"保存"')

doc.add_heading('4 AI对话使用', 1)
doc.add_paragraph('1. 点击"对话"')
doc.add_paragraph('2. 输入问题')
doc.add_paragraph('3. 发送，等待回答')

doc.add_heading('5 知识库使用', 1)
doc.add_paragraph('1. 点击"知识库"')
doc.add_paragraph('2. 点击"导入文档"')
doc.add_paragraph('3. 选择文件格式（PDF/Word/Excel/PPT/TXT）')

doc.add_heading('6 GTD任务使用', 1)
doc.add_paragraph('1. 点击"任务"')
doc.add_paragraph('2. 点击"新建任务"')
doc.add_paragraph('3. 填写任务信息')
doc.add_paragraph('4. 设置分类和优先级')

doc.add_heading('7 常见问题', 1)
doc.add_heading('7.1 启动问题', 2)
doc.add_paragraph('端口占用: 关闭占用程序或修改配置')
doc.add_paragraph('依赖缺失: 执行pip install -r requirements.txt')

doc.add_heading('7.2 使用问题', 2)
doc.add_paragraph('响应慢: 检查NPU状态，切换到BURST模式')
doc.add_paragraph('导入失败: 确认文件格式和大小')

doc.save('output/docs/用户操作手册_Antinet.docx')
print('Created: output/docs/用户操作手册_Antinet.docx')
