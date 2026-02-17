#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Create document
doc = Document()

# Title
title = doc.add_heading('Antinet 知易智能知识管家', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('技术架构文档')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('版本: 1.0.0')
doc.add_paragraph('日期: 2026-02-17')
doc.add_paragraph('平台: 骁龙X Elite AIPC')

doc.add_paragraph()

# Table of Contents
doc.add_heading('目录', 1)
toc = [
    ('1', '项目概述', '3'),
    ('2', '技术栈详解', '4'),
    ('3', '核心模块设计', '5'),
    ('4', '数据流设计', '7'),
    ('5', '部署架构', '8'),
    ('6', '安全设计', '9'),
    ('7', '性能优化', '10'),
]

for num, title, page in toc:
    doc.add_paragraph(f'{num}. {title}', style='List Number')

# Section 1: Project Overview
doc.add_heading('1 项目概述', 1)

doc.add_heading('1.1 项目简介', 2)
p = doc.add_paragraph()
p.add_run('Antinet（知易智能知识管家）是一款基于骁龙X Elite AIPC平台的端侧智能数据中枢与协同分析平台。该系统采用先进的NPU（神经网络处理器）加速技术，实现数据不出域的本地化AI推理能力，为用户提供高效、安全的智能知识管理服务。')

doc.add_heading('1.2 核心功能', 2)
features = [
    '四色卡片知识管理系统',
    'AI对话引擎',
    '八大智能体系统',
    '多模态视觉理解',
    '知识图谱构建',
    'GTD任务管理',
    '完整文档处理能力'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# Section 2: Tech Stack
doc.add_heading('2 技术栈详解', 1)

doc.add_heading('2.1 前端技术栈', 2)
p = doc.add_paragraph()
p.add_run('• 核心框架: ').bold = True
p.add_run('React 18+ 配合 TypeScript\n')
p.add_run('• 构建工具: ').bold = True
p.add_run('Vite\n')
p.add_run('• 样式方案: ').bold = True
p.add_run('TailwindCSS\n')
p.add_run('• 路由管理: ').bold = True
p.add_run('React Router 6+')

doc.add_heading('2.2 后端技术栈', 2)
p = doc.add_paragraph()
p.add_run('• Web框架: ').bold = True
p.add_run('FastAPI\n')
p.add_run('• 数据库: ').bold = True
p.add_run('SQLite + 向量数据库\n')
p.add_run('• AI/ML库: ').bold = True
p.add_run('transformers, torch\n')
p.add_run('• NPU: ').bold = True
p.add_run('QNN SDK')

# Section 3: Core Modules
doc.add_heading('3 核心模块设计', 1)

doc.add_heading('3.1 四色卡片系统', 2)
cards_table = doc.add_table(rows=5, cols=3)
cards_table.style = 'Table Grid'
headers = ['颜色', '类型', '用途']
for i, header in enumerate(headers):
    cards_table.rows[0].cells[i].text = header

card_data = [
    ('蓝色', '事实', '记录客观数据和事实信息'),
    ('绿色', '解释', '记录原因分析和解释说明'),
    ('黄色', '风险', '记录潜在风险和注意事项'),
    ('红色', '行动', '记录行动建议和待办事项')
]

for i, row_data in enumerate(card_data, 1):
    for j, cell_data in enumerate(row_data):
        cards_table.rows[i].cells[j].text = cell_data

doc.add_heading('3.2 八大智能体系统', 2)
agents = [
    '任务执行器(Task Executor) - 任务分解和执行调度',
    '事实生成器(Fact Generator) - 生成事实性内容',
    '解释器(Interpreter) - 提供分析和解释',
    '风险检测器(Risk Detector) - 识别潜在风险',
    '卡片分类器(Card Classifier) - 内容分类',
    '行动顾问(Action Advisor) - 提供行动建议',
    '记忆管理者(Memory) - 上下文和记忆管理',
    '协调器(Orchestrator) - 协调各智能体工作'
]

for agent in agents:
    doc.add_paragraph(agent, style='List Bullet')

# Section 4: Data Flow
doc.add_heading('4 数据流设计', 1)
p = doc.add_paragraph()
p.add_run('请求处理流程:\n')
p.add_run('1. 前端发起请求 → 2. API网关路由 → 3. 业务逻辑处理 → 4. 模型推理 → 5. 数据持久化 → 6. 响应返回')

# Section 5: Deployment
doc.add_heading('5 部署架构', 1)

doc.add_heading('5.1 开发环境', 2)
doc.add_paragraph('• 前端: Vite开发服务器（5173端口）')
doc.add_paragraph('• 后端: FastAPI开发服务器（8000端口）')
doc.add_paragraph('• 数据库: SQLite本地文件')

doc.add_heading('5.2 生产环境', 2)
doc.add_paragraph('• 前端: 构建为静态文件')
doc.add_paragraph('• 后端: Uvicorn多进程运行')
doc.add_paragraph('• 可选: Nginx反向代理')

doc.add_heading('5.3 端侧AIPC部署', 2)
doc.add_paragraph('• 系统直接运行在AIPC设备上')
doc.add_paragraph('• 通过NPU实现本地AI推理')
doc.add_paragraph('• 数据完全存储在本地')

# Section 6: Security
doc.add_heading('6 安全设计', 1)

doc.add_heading('6.1 数据安全', 2)
doc.add_paragraph('• 本地存储，数据不出域')
doc.add_paragraph('• 可选的加密存储')
doc.add_paragraph('• 基于角色的访问控制')

doc.add_heading('6.2 网络安全', 2)
doc.add_paragraph('• CORS配置')
doc.add_paragraph('• 输入验证')
doc.add_paragraph('• 速率限制')

# Section 7: Performance
doc.add_heading('7 性能优化', 1)

doc.add_heading('7.1 后端优化', 2)
doc.add_paragraph('• 异步处理')
doc.add_paragraph('• 连接池')
doc.add_paragraph('• 预编译SQL')
doc.add_paragraph('• 批量操作')

doc.add_heading('7.2 模型优化', 2)
doc.add_paragraph('• INT8量化推理')
doc.add_paragraph('• NPU加速')
doc.add_paragraph('• 模型缓存')

# Save
output_path = 'output/docs/技术架构文档_Antinet.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Created: {output_path}')
