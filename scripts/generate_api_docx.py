#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading('Antinet 知易智能知识管家 - API接口文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('版本: 1.0.0')
doc.add_paragraph('日期: 2026-02-17')
doc.add_paragraph('基础URL: http://localhost:8000')
doc.add_paragraph()

# Table of Contents
doc.add_heading('目录', 1)
toc_items = [
    '1. 概述',
    '2. 基础接口',
    '3. 数据管理接口',
    '4. 聊天机器人接口',
    '5. 知识管理接口',
    '6. GTD任务管理接口',
    '7. 分析接口',
    '8. 智能体接口',
    '9. 文档处理接口',
    '10. 视觉理解接口',
    '11. 技能系统接口',
    '12. 错误码说明'
]

for item in toc_items:
    doc.add_paragraph(item)

# Section 1: Overview
doc.add_heading('1 概述', 1)
doc.add_paragraph('Antinet API采用RESTful风格设计，基于FastAPI框架构建。')
doc.add_paragraph('• 基础URL: http://localhost:8000')
doc.add_paragraph('• 数据格式: JSON')
doc.add_paragraph('• 认证方式: 暂无（开发环境）')

# Section 2: Base APIs
doc.add_heading('2 基础接口', 1)

doc.add_heading('2.1 根路径', 2)
doc.add_paragraph('GET / - 获取服务基本信息')

doc.add_heading('2.2 健康检查', 2)
doc.add_paragraph('GET /api/health - 轻量级健康检查')

# Section 3: Data Management
doc.add_heading('3 数据管理接口', 1)

endpoints = [
    ('GET /api/data/team-members', '获取团队成员列表'),
    ('POST /api/data/team-members', '创建团队成员'),
    ('GET /api/data/knowledge-spaces', '获取知识空间列表'),
    ('GET /api/data/cards', '获取知识卡片列表'),
    ('POST /api/data/cards', '创建知识卡片'),
    ('POST /api/data/upload', '上传数据文件'),
]

for path, desc in endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 4: Chat
doc.add_heading('4 聊天机器人接口', 1)

chat_endpoints = [
    ('POST /api/chat', '发送聊天消息'),
    ('POST /api/chat/cards/search', '搜索知识卡片'),
]

for path, desc in chat_endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 5: Knowledge
doc.add_heading('5 知识管理接口', 1)

knowledge_endpoints = [
    ('GET /api/knowledge/search', '搜索知识库'),
    ('GET /api/knowledge/graph', '获取知识图谱'),
]

for path, desc in knowledge_endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 6: GTD
doc.add_heading('6 GTD任务管理接口', 1)

gtd_endpoints = [
    ('GET /api/data/gtd/tasks', '获取所有GTD任务'),
    ('POST /api/data/gtd/tasks', '创建GTD任务'),
    ('PUT /api/data/gtd/tasks/{task_id}', '更新GTD任务'),
    ('DELETE /api/data/gtd/tasks/{task_id}', '删除GTD任务'),
]

for path, desc in gtd_endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 7: Analysis
doc.add_heading('7 分析接口', 1)

analysis_endpoints = [
    ('POST /api/analyze', '数据分析接口'),
    ('GET /api/performance/benchmark', '性能基准测试'),
]

for path, desc in analysis_endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 8: Agent
doc.add_heading('8 智能体接口', 1)

agent_endpoints = [
    ('POST /api/agent/execute', '执行智能体任务'),
    ('GET /api/agent/status', '获取智能体状态'),
]

for path, desc in agent_endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 9: Document
doc.add_heading('9 文档处理接口', 1)

doc_endpoints = [
    ('POST /api/pdf/analyze', 'PDF分析'),
    ('POST /api/excel/analyze', 'Excel分析'),
    ('POST /api/ppt/generate', 'PPT生成'),
]

for path, desc in doc_endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 10: Vision
doc.add_heading('10 视觉理解接口', 1)
doc.add_paragraph('POST /api/vision/analyze - 图像分析', style='List Bullet')

# Section 11: Skills
doc.add_heading('11 技能系统接口', 1)

skill_endpoints = [
    ('POST /api/skill/execute', '执行技能'),
    ('GET /api/skills', '获取技能列表'),
]

for path, desc in skill_endpoints:
    doc.add_paragraph(f'{path} - {desc}', style='List Bullet')

# Section 12: Error Codes
doc.add_heading('12 错误码说明', 1)

error_table = doc.add_table(rows=6, cols=2)
error_table.style = 'Table Grid'
error_table.rows[0].cells[0].text = '错误码'
error_table.rows[0].cells[1].text = '说明'

errors = [
    ('200', '成功'),
    ('400', '请求参数错误'),
    ('404', '资源不存在'),
    ('500', '服务器内部错误'),
    ('503', '服务不可用'),
]

for i, (code, desc) in enumerate(errors, 1):
    error_table.rows[i].cells[0].text = code
    error_table.rows[i].cells[1].text = desc

# Save
output_path = 'output/docs/API接口文档_Antinet.docx'
doc.save(output_path)
print(f'Created: {output_path}')
