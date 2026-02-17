#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Antinet 知易智能知识管家 - 部署指南', 0)
doc.add_paragraph('版本: 1.0.0')
doc.add_paragraph('日期: 2026-02-17')
doc.add_paragraph()

doc.add_heading('1 部署概述', 1)
doc.add_heading('1.1 部署模式', 2)
doc.add_paragraph('Antinet支持三种部署模式：')
doc.add_paragraph('1. 开发环境 - 适合调试', style='List Bullet')
doc.add_paragraph('2. 生产环境 - 适合正式上线', style='List Bullet')
doc.add_paragraph('3. 端侧AIPC部署 - 核心场景，数据不出域', style='List Bullet')

doc.add_heading('2 环境准备', 1)
doc.add_heading('2.1 硬件要求', 2)
doc.add_paragraph('开发环境: CPU Intel Core i5+, 内存 8GB+, 存储 20GB+')
doc.add_paragraph('端侧AIPC: 骁龙X Elite, 内存 16GB+, 存储 50GB+')

doc.add_heading('2.2 软件要求', 2)
doc.add_paragraph('操作系统: Windows 10/11 64位')
doc.add_paragraph('Python: 3.9+')
doc.add_paragraph('Node.js: 18+')

doc.add_heading('2.3 依赖安装', 2)
doc.add_paragraph('# Python依赖')
doc.add_paragraph('cd backend')
doc.add_paragraph('python -m venv venv')
doc.add_paragraph('venv\\Scripts\\activate')
doc.add_paragraph('pip install -r requirements.txt')
doc.add_paragraph()
doc.add_paragraph('# Node.js依赖')
doc.add_paragraph('npm install')

doc.add_heading('3 开发环境部署', 1)
doc.add_heading('3.1 启动后端', 2)
doc.add_paragraph('cd backend')
doc.add_paragraph('python main.py')
doc.add_paragraph('服务运行在 http://localhost:8000')

doc.add_heading('3.2 启动前端', 2)
doc.add_paragraph('cd src')
doc.add_paragraph('npm run dev')
doc.add_paragraph('服务运行在 http://localhost:5173')

doc.add_heading('4 生产环境部署', 1)
doc.add_heading('4.1 前端构建', 2)
doc.add_paragraph('cd src')
doc.add_paragraph('npm run build')

doc.add_heading('4.2 后端配置', 2)
doc.add_paragraph('uvicorn main:app --host 0.0.0.0 --port 8000 --workers 4')

doc.add_heading('5 端侧AIPC部署', 1)
doc.add_heading('5.1 NPU环境配置', 2)
doc.add_paragraph('1. 安装QNN SDK')
doc.add_paragraph('2. 配置环境变量')
doc.add_paragraph('3. 放置模型文件')

doc.add_heading('5.2 启动验证', 2)
doc.add_paragraph('curl http://localhost:8000/api/health')

doc.add_heading('6 数据管理', 1)
doc.add_paragraph('data/ - 主数据库')
doc.add_paragraph('data/uploads/ - 上传文件')
doc.add_paragraph('data/exports/ - 导出文件')
doc.add_paragraph('data/backups/ - 备份')

doc.add_heading('7 常见问题', 1)
doc.add_heading('7.1 启动问题', 2)
doc.add_paragraph('端口占用: 关闭占用程序或修改端口')
doc.add_paragraph('依赖缺失: 重新安装依赖')

doc.add_heading('7.2 使用问题', 2)
doc.add_paragraph('响应慢: 检查NPU，切换性能模式')
doc.add_paragraph('导入失败: 检查文件格式和大小')

doc.save('output/docs/部署指南_Antinet.docx')
print('Created: output/docs/部署指南_Antinet.docx')
