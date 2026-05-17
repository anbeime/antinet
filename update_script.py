from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('/mnt/c/D/zhiyi/路演稿_决赛答辩.docx')

insert_content = [
    ("", "normal"),
    ("核心功能与现场演示路径", "heading"),
    ("\u23f1 建议时长：1分钟", "normal"),
    ("", "normal"),
    ("[PPT: 第7页] 核心功能与现场演示路径", "normal"),
    ("", "normal"),
    ("各位评委老师，接下来的演示环节，我会带大家一步步体验Antinet的核心功能。", "normal"),
    ("", "normal"),
    ("步骤1：启动系统", "bold"),
    ("双击 start.bat 一键启动，系统将自动启动后端 FastAPI（端口8000）和前端 Vite（端口3000），并在10秒内将Qwen2-7B-SSD模型加载到NPU。", "normal"),
    ("", "normal"),
    ("步骤2：进入主界面", "bold"),
    ("在浏览器打开 http://localhost:3000，即可看到知易主界面——包含上传区、卡片展示区和Agent状态面板。演示时，我会向各位确认：8个Agent全部在线，状态指示灯均为绿色。", "normal"),
    ("", "normal"),
    ("步骤3：上传文档并演示完整流程", "bold"),
    ("我会向大家展示一条完整的主线流程：上传一份PDF文档 → 系统自动解析 → 输出四色卡片。每一步的输入、处理、输出都会清晰展示，突出NPU加速效果——响应时间小于500毫秒。", "normal"),
    ("", "normal"),
    ("步骤4：补充展示（如时间允许）", "bold"),
    ("如果时间充裕，我还会演示：知识库搜索、卡片关联追溯、历史记录回溯等扩展功能。", "normal"),
    ("", "normal"),
    ("特别说明：本系统完全离线运行，无需联网、无需账号、无需特殊外设，真正实现了'数据不出域'的隐私保护。演示设备为骁龙X Elite AIPC，备有CPU后备推理模式和录屏视频双重保险。", "normal"),
    ("", "normal"),
]

# 找到"四、技术架构"前面的段落
insert_at = None
for i, para in enumerate(doc.paragraphs):
    if "四、技术架构" in para.text:
        insert_at = i
        break

if insert_at is None:
    print("ERROR: 找不到插入点")
    exit(1)

print(f"在段落 {insert_at} 处插入内容")

for item_text, style in reversed(insert_content):
    p = doc.paragraphs[insert_at].insert_paragraph_before()
    if style == "heading":
        run = p.add_run(item_text)
        run.bold = True
        run.font.size = Pt(14)
    elif style == "bold":
        run = p.add_run(item_text)
        run.bold = True
    else:
        p.add_run(item_text)

doc.save('/mnt/c/D/zhiyi/路演稿_决赛答辩.docx')
print("保存成功！")

# 验证
doc2 = Document('/mnt/c/D/zhiyi/路演稿_决赛答辩.docx')
for i, para in enumerate(doc2.paragraphs):
    if "核心功能与现场演示路径" in para.text:
        print(f"验证成功：第 {i} 段找到目标内容")
        break